# -*- coding: utf-8 -*-
"""
Created on Sat Oct  3 20:31:58 2020

@author: wenlei
"""
from docx import Document
from docx.shared import RGBColor#设置字体
class Author():
    def __init__(self,FamilyName,MiddleName,LastName):
        self.FamilyName=FamilyName.title()
        self.MiddleName=MiddleName.title()
        self.LastName=LastName.title()
        self.Name=[]
        self.update()
        self.FamilyName=FamilyName.upper()
        self.MiddleName=MiddleName.upper()
        self.LastName=LastName.upper()
        self.update()
        self.FamilyName=FamilyName.lower()
        self.MiddleName=MiddleName.lower()
        self.LastName=LastName.lower()
        self.update()
        self.FamilyName=FamilyName.title()
        self.MiddleName=MiddleName.title()
        self.LastName=LastName.lower()
        self.update()
        print(self.Name)
    def update(self):
        if self.MiddleName=="":
            self.Name.append("%s, %s"%(self.FamilyName,self.LastName))
            self.Name.append("%s, %s."%(self.FamilyName,self.LastName[0]))
            self.Name.append("%s %s"%(self.LastName,self.FamilyName,))
            self.Name.append("%s, %s."%(self.FamilyName,self.LastName[0]))
            self.Name.append("%s %s"%(self.FamilyName,self.LastName))
            self.Name.append("%s. %s"%(self.FamilyName,self.LastName[0]))
            self.Name.append("%s. %s"%(self.LastName[0],self.FamilyName))
        else:
            self.Name.append("%s, %s%s"%(self.FamilyName,self.MiddleName,self.LastName))
            self.Name.append("%s, %s.%s."%(self.FamilyName,self.MiddleName[0],self.LastName[0]))
            self.Name.append("%s%s %s"%(self.MiddleName,self.LastName,self.FamilyName))
            self.Name.append("%s, %s. %s."%(self.FamilyName,self.MiddleName[0],self.LastName[0]))
            self.Name.append("%s, %s."%(self.FamilyName,self.MiddleName[0]))
            self.Name.append("%s %s%s"%(self.FamilyName,self.MiddleName,self.LastName))
            self.Name.append("%s. %s%s"%(self.FamilyName,self.MiddleName[0],self.LastName[0]))
            self.Name.append("%s%s. %s"%(self.MiddleName[0],self.LastName[0],self.FamilyName))
            self.Name.append("%s %s-%s"%(self.FamilyName,self.MiddleName,self.LastName))
            self.Name.append("%s, %s%s"%(self.FamilyName,self.MiddleName[0],self.LastName[0]))
             
        
docfile="代表性论文及检索列表1003v2.docx"
ipaper=7
#OrgTitle="Variable order fractional differential operators in anomalous diffusion modeling"
#OrgAuthor=[Author('Sun',"Hong",'Guang'),Author("Chen","","Wen"),Author("Chen","Yang","Quan")]
OrgTitle="Impact of climate change on flood and drought events in Huaihe River Basin, China"
OrgAuthor=[Author('Yang',"Chuan",'Guo'),Author("Yu","Zhong","Bo"),Author("Hao","Zhen","Chun"),Author("Zhang","Jiang","Yun"),Author("Zhu","Jian","Ting")]
f=open('savedrecs (1).txt','r',encoding='utf-8')
d=f.readlines()
a=[]
author=[]
title=[]
date=[]
jour=[]
for i in range(1,len(d)):
    a.append(d[i].split("\t"))
    if a[i-1][1]=="":
        author.append(a[i-1][8])
    else:
        author.append(a[i-1][1])
    title.append(a[i-1][9])
    if a[i-1][32]=="":
        date.append(a[i-1][33])
    else:
        date.append(a[i-1][32])
    jour.append(a[i-1][17])
doc=Document()

doc.add_heading(OrgTitle)
zuozhe="作者:"
for l in OrgAuthor:
    zuozhe+=" %s;"%l.Name[0]
doc.add_paragraph(zuozhe[:-1])

table = doc.add_table(rows=1,cols=5,style="Medium Grid 1 Accent 1")
hd_cells=table.rows[0].cells
hd_cells[0].text='序号'
hd_cells[1].text='论文名称'
hd_cells[2].text='作者'
hd_cells[3].text='期刊'
hd_cells[4].text='日期'
ziyin=0
for i in range(len(author)):
    row_cells=table.add_row().cells
    row_cells[0].text=str(i+1)
    row_cells[1].text=title[i]
    my_cell=row_cells[2]
    my_paragraph=my_cell.paragraphs[0]
    my_author=author[i].split(";")
    run=[]
    for ator in my_author:
        wl=0
        for k in OrgAuthor:
            for name in k.Name:
                if name in ator:
                    wl=1
        if wl==1:
            run.append(my_paragraph.add_run(ator))
        else:
            my_paragraph.add_run(ator)
    if len(run)>0:
        ziyin+=1
        red=RGBColor(255,0,0)
        for j in run:
            j.font.color.rgb=red
    row_cells[3].text=jour[i]
    row_cells[4].text=date[i]
doc.save('%i.docx'%ipaper)
result="检索被引文献: %i;他引次数: %i;自引次数为: %i "%(len(author),len(author)-ziyin,ziyin)
#open('result.txt','w').writelines(result)
print(result)

docwt=Document(docfile)
table1=docwt.tables[0]
cl=table1.cell(ipaper,7)
cl.text=str(len(author))

cl=table1.cell(ipaper,8)
cl.text=str(len(author)-ziyin)
docwt.save(docfile)